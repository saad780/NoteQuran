surah = 2
start = 60
stop = 101

import requests
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time
from joblib import Parallel, delayed
import multiprocessing

url = 'https://api.alquran.cloud/ayah/' + str(surah) + ':'

document = Document()
section = document.sections[0]
section.left_margin = Inches(.25)
section.right_margin = Inches(.25)
section.top_margin = Inches(.17)
section.bottom_margin = Inches(.17)

table = document.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

s = requests.session()
start_time = time.time()

ncores = multiprocessing.cpu_count()
inputs = range(start, stop+1)

def grabAyah(x):
  print(x)
  while (True):
    try:
      ayah = s.get(url + str(x)).json()['data']['text']
      tran = s.get(url + str(x) + '/en.sahih').json()['data']['text']
      break
    except:
      print("except")
      time.sleep(.3)  # Request fails so pause needed before retry.
  return(ayah, tran)

results =  Parallel(n_jobs=ncores)(delayed(grabAyah)(i) for i in inputs)

x=start
for ayah,tran in results:
  row_cells = table.add_row().cells
  row_cells[0].text = str(x) + ". " + tran
  row_cells[1].text = ayah
  row_cells[1].paragraphs[0].runs[0].font.name = 'Cambria'
  row_cells[1].paragraphs[0].runs[0].font.size = Pt(18)
  row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
  table.add_row()
  x+=1

document.save('test.docx')

print(time.time()-start_time)